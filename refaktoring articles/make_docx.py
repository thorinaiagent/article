#!/usr/bin/env python3
"""
make_docx.py — собирает DOCX из отредактированной статьи (Markdown).

Использование:
    python3 make_docx.py <входной_файл.md> [название_статьи]

Что делает:
    - читает Markdown-файл с отредактированной статьёй;
    - конвертирует базовую разметку (заголовки, списки, жирный, код) в DOCX;
    - сохраняет результат в папку ~/Desktop/Отредактированные статьи/;
    - печатает полный путь к готовому файлу.

Зависимость: python-docx. Если не установлена:
    pip3 install python-docx
"""

import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
except ImportError:
    sys.exit(
        "Не установлена библиотека python-docx.\n"
        "Установи одной командой:  pip3 install python-docx"
    )

# Папка на рабочем столе, куда складываем все отредактированные статьи
OUTPUT_DIR = os.path.expanduser("~/Desktop/Отредактированные статьи")


def add_runs_with_inline(paragraph, text):
    """Разбирает инлайновый **жирный** и `код`, добавляя runs в абзац."""
    # Разбиваем по **...** и `...`, сохраняя разделители
    tokens = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(tok)


def markdown_to_docx(md_text, title=None):
    doc = Document()

    # Базовый шрифт
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    if title:
        doc.add_heading(title, level=0)

    lines = md_text.splitlines()
    in_code_block = False
    code_buffer = []

    for line in lines:
        stripped = line.strip()

        # Блоки кода ```...```
        if stripped.startswith("```"):
            if in_code_block:
                # закрываем блок
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buffer))
                run.font.name = "Courier New"
                run.font.size = Pt(10)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_buffer.append(line)
            continue

        # Пустая строка
        if not stripped:
            continue

        # Горизонтальный разделитель
        if re.fullmatch(r"-{3,}", stripped):
            continue

        # Заголовки
        m = re.match(r"^(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=min(level, 4))
            continue

        # Нумерованный список
        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline(p, m.group(1))
            continue

        # Маркированный список
        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_inline(p, m.group(1))
            continue

        # Обычный абзац
        p = doc.add_paragraph()
        add_runs_with_inline(p, stripped)

    return doc


def main():
    if len(sys.argv) < 2:
        sys.exit("Укажи входной файл: python3 make_docx.py <файл.md> [название]")

    input_path = os.path.expanduser(sys.argv[1])
    if not os.path.isfile(input_path):
        sys.exit(f"Файл не найден: {input_path}")

    with open(input_path, encoding="utf-8") as f:
        md_text = f.read()

    # Название: из аргумента, либо из первого заголовка, либо из имени файла
    if len(sys.argv) >= 3:
        title = sys.argv[2]
    else:
        first_heading = next(
            (l.lstrip("# ").strip() for l in md_text.splitlines() if l.startswith("#")),
            None,
        )
        title = first_heading or os.path.splitext(os.path.basename(input_path))[0]

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Безопасное имя файла + дата, чтобы версии не перетирали друг друга
    safe_title = re.sub(r"[^\w\s-]", "", title).strip()[:80] or "статья"
    stamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    out_path = os.path.join(OUTPUT_DIR, f"{safe_title} — отредактировано {stamp}.docx")

    doc = markdown_to_docx(md_text, title=title)
    doc.save(out_path)

    print(out_path)


if __name__ == "__main__":
    main()
